#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
用户需求研究报告生成器
生成Word格式的专业研究报告
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import json
import sys
from datetime import datetime

class UserNeedsReportGenerator:
    def __init__(self, data_file=None):
        self.doc = Document()
        self.data = self.load_data(data_file) if data_file else self.get_default_data()
        self.setup_styles()
    
    def load_data(self, data_file):
        """加载研究数据"""
        try:
            with open(data_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return self.get_default_data()
    
    def get_default_data(self):
        """默认数据结构"""
        return {
            "project_name": "用户需求研究",
            "report_date": datetime.now().strftime("%Y-%m-%d"),
            "version": "1.0",
            "team": "产品研究团队",
            
            "overview": {
                "background": "",
                "objectives": [],
                "methods": [],
                "sample": ""
            },
            
            "hierarchy": {
                "current_level": "L1-L2",
                "unmet_levels": ["L3", "L4", "L5"],
                "opportunities": []
            },
            
            "pain_points": [],
            
            "user_segments": [],
            
            "decision_journey": [],
            
            "motivations": [],
            
            "behavior_logic": [],
            
            "insights": [],
            "recommendations": [],
            "action_plan": []
        }
    
    def setup_styles(self):
        """设置文档样式"""
        # 设置中文字体
        self.doc.styles['Normal'].font.name = 'Microsoft YaHei'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        self.doc.styles['Normal'].font.size = Pt(10.5)
    
    def add_heading(self, text, level=1):
        """添加标题"""
        heading = self.doc.add_heading(text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in heading.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            if level == 1:
                run.font.color.rgb = RGBColor(0, 51, 102)
                run.font.size = Pt(18)
            elif level == 2:
                run.font.color.rgb = RGBColor(0, 102, 153)
                run.font.size = Pt(14)
        return heading
    
    def add_paragraph(self, text, bold=False, color=None):
        """添加段落"""
        p = self.doc.add_paragraph(text)
        for run in p.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.bold = bold
            if color:
                run.font.color.rgb = color
        return p
    
    def add_cover(self):
        """添加封面"""
        # 空行
        for _ in range(6):
            self.doc.add_paragraph()
        
        # 标题
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("用户需求研究报告")
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        
        # 空行
        for _ in range(4):
            self.doc.add_paragraph()
        
        # 项目信息
        info_items = [
            ("项目名称", self.data.get("project_name", "")),
            ("报告日期", self.data.get("report_date", "")),
            ("版本号", self.data.get("version", "")),
            ("研究团队", self.data.get("team", ""))
        ]
        
        for label, value in info_items:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{label}：{value}")
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(12)
        
        # 分页
        self.doc.add_page_break()
    
    def add_overview(self):
        """添加研究概述"""
        self.add_heading("第一章：研究概述", level=1)
        
        # 研究背景
        self.add_heading("1.1 研究背景", level=2)
        bg = self.data.get("overview", {}).get("background", "")
        self.add_paragraph(bg if bg else "（请填写研究背景）")
        
        # 研究目标
        self.add_heading("1.2 研究目标", level=2)
        objectives = self.data.get("overview", {}).get("objectives", [])
        if objectives:
            for i, obj in enumerate(objectives, 1):
                self.add_paragraph(f"{i}. {obj}")
        else:
            self.add_paragraph("1. 深入了解目标用户需求层级")
            self.add_paragraph("2. 挖掘用户核心痛点")
            self.add_paragraph("3. 识别不同用户群体特征")
            self.add_paragraph("4. 分析用户决策链路")
            self.add_paragraph("5. 揭示消费动机和行为逻辑")
        
        # 研究方法
        self.add_heading("1.3 研究方法", level=2)
        methods = self.data.get("overview", {}).get("methods", [])
        if methods:
            for method in methods:
                self.add_paragraph(f"• {method}")
        else:
            self.add_paragraph("• 用户深度访谈")
            self.add_paragraph("• 问卷调查")
            self.add_paragraph("• 行为数据分析")
            self.add_paragraph("• 竞品对比研究")
        
        # 样本说明
        self.add_heading("1.4 样本说明", level=2)
        sample = self.data.get("overview", {}).get("sample", "")
        self.add_paragraph(sample if sample else "（请填写样本说明）")
    
    def add_hierarchy(self):
        """添加需求层级分析"""
        self.add_heading("第二章：需求层级分析", level=1)
        
        self.add_heading("2.1 需求金字塔", level=2)
        self.add_paragraph("基于马斯洛需求层次理论，将用户需求分为五个层级：")
        
        hierarchy_table = self.doc.add_table(rows=6, cols=3)
        hierarchy_table.style = 'Light Grid Accent 1'
        
        # 表头
        headers = ['层级', '需求类型', '产品体现']
        for i, header in enumerate(headers):
            cell = hierarchy_table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # 数据
        levels = [
            ('L5', '自我实现', '成长、意义、成就'),
            ('L4', '尊重需求', '认可、地位、掌控'),
            ('L3', '社交需求', '归属、连接、认同'),
            ('L2', '安全需求', '稳定、保障、信任'),
            ('L1', '生理需求', '功能、效率、基础')
        ]
        
        for i, (level, need_type, product) in enumerate(levels, 1):
            row = hierarchy_table.rows[i]
            row.cells[0].text = level
            row.cells[1].text = need_type
            row.cells[2].text = product
        
        self.add_heading("2.2 当前满足度评估", level=2)
        current = self.data.get("hierarchy", {}).get("current_level", "L1-L2")
        self.add_paragraph(f"当前产品主要满足：{current}")
        
        self.add_heading("2.3 未满足需求识别", level=2)
        unmet = self.data.get("hierarchy", {}).get("unmet_levels", ["L3", "L4", "L5"])
        self.add_paragraph(f"未满足的需求层级：{', '.join(unmet)}")
        
        self.add_heading("2.4 机会点总结", level=2)
        opportunities = self.data.get("hierarchy", {}).get("opportunities", [])
        if opportunities:
            for opp in opportunities:
                self.add_paragraph(f"• {opp}")
        else:
            self.add_paragraph("• 高层级需求是差异化机会")
            self.add_paragraph("• 跨层级跃迁可实现创新突破")
    
    def add_pain_points(self):
        """添加痛点挖掘"""
        self.add_heading("第三章：痛点挖掘", level=1)
        
        self.add_heading("3.1 痛点清单", level=2)
        pain_points = self.data.get("pain_points", [])
        if pain_points:
            for i, pp in enumerate(pain_points, 1):
                self.add_paragraph(f"{i}. {pp.get('description', '')}")
                self.add_paragraph(f"   强度：{pp.get('intensity', '')} | 频率：{pp.get('frequency', '')} | 优先级：{pp.get('priority', '')}")
        else:
            self.add_paragraph("（请根据研究数据填写痛点清单）")
        
        self.add_heading("3.2 痛点优先级", level=2)
        self.add_paragraph("评分标准：强度(1-5) × 2 + 频率(1-5) × 1.5 + 紧迫性(1-5) × 1.5 + 支付意愿(1-5) × 1")
        
        priority_table = self.doc.add_table(rows=5, cols=3)
        priority_table.style = 'Light Grid Accent 1'
        
        headers = ['得分区间', '优先级', '行动建议']
        for i, header in enumerate(headers):
            cell = priority_table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        priorities = [
            ('20-25', 'P0', '必须解决，核心功能'),
            ('15-19', 'P1', '重要痛点，尽快解决'),
            ('10-14', 'P2', '一般痛点，排期解决'),
            ('<10', 'P3', '观察，暂不处理')
        ]
        
        for i, (score, priority, action) in enumerate(priorities, 1):
            row = priority_table.rows[i]
            row.cells[0].text = score
            row.cells[1].text = priority
            row.cells[2].text = action
        
        self.add_heading("3.3 根因分析", level=2)
        self.add_paragraph("（使用5Why分析法进行根因分析）")
        
        self.add_heading("3.4 解决方向", level=2)
        self.add_paragraph("（基于痛点优先级提出解决方向）")
    
    def add_user_segments(self):
        """添加用户分层"""
        self.add_heading("第四章：用户分层", level=1)
        
        self.add_heading("4.1 分层维度", level=2)
        self.add_paragraph("分层维度包括：")
        self.add_paragraph("• 基础属性：年龄、性别、收入、地域、职业")
        self.add_paragraph("• 行为属性：使用频率、使用深度、功能偏好")
        self.add_paragraph("• 价值属性：付费能力、生命周期、传播价值")
        
        self.add_heading("4.2 用户画像", level=2)
        segments = self.data.get("user_segments", [])
        if segments:
            for segment in segments:
                self.add_heading(segment.get('name', ''), level=3)
                self.add_paragraph(f"占比：{segment.get('percentage', '')}")
                self.add_paragraph(f"特征：{segment.get('characteristics', '')}")
                self.add_paragraph(f"策略：{segment.get('strategy', '')}")
        else:
            self.add_paragraph("（请根据研究数据填写用户画像）")
        
        self.add_heading("4.3 分层策略", level=2)
        self.add_paragraph("（针对不同用户群体的产品策略）")
        
        self.add_heading("4.4 优先级排序", level=2)
        self.add_paragraph("（确定核心目标用户群体）")
    
    def add_decision_journey(self):
        """添加决策链路"""
        self.add_heading("第五章：决策链路", level=1)
        
        self.add_heading("5.1 决策旅程", level=2)
        self.add_paragraph("用户决策阶段：认知 → 兴趣 → 评估 → 决策 → 购买 → 使用 → 复购/流失")
        
        journey_table = self.doc.add_table(rows=8, cols=2)
        journey_table.style = 'Light Grid Accent 1'
        
        headers = ['阶段', '关键问题']
        for i, header in enumerate(headers):
            cell = journey_table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        stages = [
            ('认知', '用户如何知道我们？'),
            ('兴趣', '什么吸引了用户？'),
            ('评估', '用户对比什么？'),
            ('决策', '临门一脚是什么？'),
            ('购买', '购买过程顺畅吗？'),
            ('使用', '用户如何使用？'),
            ('复购', '为什么回来/离开？')
        ]
        
        for i, (stage, question) in enumerate(stages, 1):
            row = journey_table.rows[i]
            row.cells[0].text = stage
            row.cells[1].text = question
        
        self.add_heading("5.2 关键触点", level=2)
        self.add_paragraph("（识别影响用户决策的关键触点）")
        
        self.add_heading("5.3 决策因素", level=2)
        self.add_paragraph("（分析用户的决策因素权重）")
        
        self.add_heading("5.4 优化建议", level=2)
        self.add_paragraph("（基于决策链路提出优化建议）")
    
    def add_motivations(self):
        """添加消费动机"""
        self.add_heading("第六章：消费动机", level=1)
        
        self.add_heading("6.1 动机类型", level=2)
        self.add_paragraph("动机分为三类：")
        self.add_paragraph("• 功能动机（理性）：解决问题、提升效率、节省成本")
        self.add_paragraph("• 情感动机（感性）：获得愉悦、减少焦虑、自我表达")
        self.add_paragraph("• 社会动机（关系）：社交连接、地位彰显、群体归属")
        
        self.add_heading("6.2 动机强度", level=2)
        motivations = self.data.get("motivations", [])
        if motivations:
            for m in motivations:
                self.add_paragraph(f"• {m.get('type', '')}：{m.get('strength', '')}")
        else:
            self.add_paragraph("（请根据研究数据填写动机强度）")
        
        self.add_heading("6.3 动机-产品映射", level=2)
        self.add_paragraph("（将用户动机映射到产品功能）")
        
        self.add_heading("6.4 价值主张优化", level=2)
        self.add_paragraph("（基于动机分析优化价值主张）")
    
    def add_behavior_logic(self):
        """添加行为逻辑"""
        self.add_heading("第七章：行为逻辑", level=1)
        
        self.add_heading("7.1 用户旅程", level=2)
        self.add_paragraph("基于BJ Fogg模型：行为 = 动机 × 能力 × 提示")
        
        self.add_heading("7.2 行为模型", level=2)
        self.add_paragraph("• 动机：用户多想要？")
        self.add_paragraph("• 能力：用户做得到吗？")
        self.add_paragraph("• 提示：用户记得吗？")
        
        self.add_heading("7.3 关键行为数据", level=2)
        behaviors = self.data.get("behavior_logic", [])
        if behaviors:
            for b in behaviors:
                self.add_paragraph(f"• {b.get('metric', '')}：{b.get('value', '')}")
        else:
            self.add_paragraph("• 活跃度：DAU/MAU、使用时长、频次")
            self.add_paragraph("• 留存：次日/7日/30日留存")
            self.add_paragraph("• 转化：注册率、付费率、复购率")
            self.add_paragraph("• 传播：分享率、推荐率、NPS")
        
        self.add_heading("7.4 行为优化点", level=2)
        self.add_paragraph("（基于行为分析提出优化点）")
    
    def add_insights(self):
        """添加核心洞察与建议"""
        self.add_heading("第八章：核心洞察与建议", level=1)
        
        self.add_heading("8.1 关键发现", level=2)
        insights = self.data.get("insights", [])
        if insights:
            for i, insight in enumerate(insights, 1):
                self.add_paragraph(f"{i}. {insight}")
        else:
            self.add_paragraph("（请总结关键发现）")
        
        self.add_heading("8.2 战略建议", level=2)
        recommendations = self.data.get("recommendations", [])
        if recommendations:
            for i, rec in enumerate(recommendations, 1):
                self.add_paragraph(f"{i}. {rec}")
        else:
            self.add_paragraph("（请提出战略建议）")
        
        self.add_heading("8.3 行动计划", level=2)
        action_plan = self.data.get("action_plan", [])
        if action_plan:
            for item in action_plan:
                self.add_paragraph(f"• {item}")
        else:
            self.add_paragraph("（请制定行动计划）")
        
        self.add_heading("8.4 风险提醒", level=2)
        self.add_paragraph("• 研究样本局限性")
        self.add_paragraph("• 市场环境变化")
        self.add_paragraph("• 竞品动态影响")
    
    def generate(self, output_path=None):
        """生成完整报告"""
        # 封面
        self.add_cover()
        
        # 目录页
        self.add_heading("目录", level=1)
        toc_items = [
            "第一章：研究概述",
            "第二章：需求层级分析",
            "第三章：痛点挖掘",
            "第四章：用户分层",
            "第五章：决策链路",
            "第六章：消费动机",
            "第七章：行为逻辑",
            "第八章：核心洞察与建议"
        ]
        for item in toc_items:
            self.add_paragraph(item)
        
        # 各章节
        self.add_overview()
        self.add_hierarchy()
        self.add_pain_points()
        self.add_user_segments()
        self.add_decision_journey()
        self.add_motivations()
        self.add_behavior_logic()
        self.add_insights()
        
        # 保存
        if not output_path:
            output_path = f"用户需求研究报告_{datetime.now().strftime('%Y%m%d')}.docx"
        
        self.doc.save(output_path)
        print(f"报告已生成：{output_path}")
        return output_path


def main():
    """主函数"""
    import argparse
    
    parser = argparse.ArgumentParser(description='生成用户需求研究报告')
    parser.add_argument('--data', '-d', help='数据文件路径(JSON格式)')
    parser.add_argument('--output', '-o', help='输出文件路径')
    
    args = parser.parse_args()
    
    generator = UserNeedsReportGenerator(args.data)
    output_path = generator.generate(args.output)
    
    print(f"✅ 报告生成完成：{output_path}")


if __name__ == '__main__':
    main()
